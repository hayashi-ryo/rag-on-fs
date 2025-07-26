import os
from docx import Document
from openpyxl import Workbook
from fpdf import FPDF

# 作成先ディレクトリ
base_dir = "samples"
os.makedirs(f"{base_dir}/docx", exist_ok=True)
os.makedirs(f"{base_dir}/xlsx", exist_ok=True)
os.makedirs(f"{base_dir}/pdf", exist_ok=True)
os.makedirs(f"{base_dir}/txt", exist_ok=True)

# Wordファイル（.docx）
def create_docx(path):
    doc = Document()
    doc.add_heading('設計仕様書サンプル', 0)
    doc.add_paragraph('この文書はPoC用に自動生成されたサンプルです。')
    doc.add_paragraph('第1章：概要\n第2章：機能一覧\n第3章：補足事項')
    doc.save(path)

# Excelファイル（.xlsx）
def create_xlsx(path):
    wb = Workbook()
    ws = wb.active
    ws.title = "サンプルデータ"
    ws.append(["項目", "数値", "備考"])
    ws.append(["応答時間", 123, "ms"])
    ws.append(["CPU使用率", 45, "%"])
    wb.save(path)

# PDFファイル（英文に制限）
def create_pdf(path):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Sample Manual PDF", ln=True)
    pdf.multi_cell(0, 10, txt="This is a sample PDF for testing.\nIt contains plain English text.\nEach page is processed individually.")
    pdf.output(path)

# テキストファイル（.txt）
def create_txt(path):
    with open(path, "w", encoding="utf-8") as f:
        f.write("これはサンプルのテキストファイルです。\nメモや自由記述に使われることを想定しています。")

# ファイル生成呼び出し
create_docx(f"{base_dir}/docx/sample_doc.docx")
create_xlsx(f"{base_dir}/xlsx/sample_sheet.xlsx")
create_pdf(f"{base_dir}/pdf/sample_manual.pdf")
create_txt(f"{base_dir}/txt/sample_note.txt")

print("✅ サンプル文書の生成が完了しました。")
